'''
Code written and developed by Jan Van Zeghbroeck
https://github.com/janvanzeghbroeck
'''


import docx
import os
import pandas as pd
import pickle
import re
import sys
import numpy as np

# check the python version
if sys.version_info[0] == 3:
    version_folder = 'data3'
else:
    version_folder = 'data'# current_folder = os.getcwd()

def process_table(table):
    n_rows = len(table.rows) #19
    # n_cols = len(table.columns)#6

    # creates a table
    table_ = []
    for i in range(n_rows):
        row = table.rows[i]
        row_ = []
        for cell in row.cells:
            row_.append(cell.text)

        table_.append(row_)

    # creates data frame
    df = pd.DataFrame(table_)
    df.columns = df.iloc[0,:].values #rename columns
    df.drop(0,axis = 0,inplace = True) #delete column name row
    return df

class ReadNewBeers(object):

    def __init__(self,cold_start = False):
        self.cold_start = cold_start
        self.wanted_comments = [ 'aroma', 'mouthfeel'] #'Visual Comment  ',

    def Update_inter_file(self, file_path, data):
        if not self.cold_start:
            pickle_file = list(pickle.load(open(file_path,'rb')))
            data = pickle_file + data

        pickle.dump( data, open(file_path, "wb" ) )
        return data

    def Get_files_from_folders(self,folder):
        # get only folders not files.
        sub_folders = [fold for fold in os.listdir(folder) if '.' not in fold]

        self.file_names = [fold + '/' + file_name for fold in sub_folders for file_name in os.listdir(folder + '/' + fold) if file_name.split('.')[-1] == 'docx']

    def Read_new_beer_files(self,folder,from_list = False):

        '''
        reads all the tables from docx files and puts it into a list of lists of datafrems
        beers(file) --> beer --> tables --> table
        '''

        # file_name = 'Cherry Brown FTR Sep 16 15.docx'
        beers = []
        beer_names = []
        ifiles = []

        if self.cold_start:
            finished_files = []
            print('!!! COLD START !!!')
        else:# load what files we have alredy looked at
            finished_files = list( pickle.load( open('{}/inter_files/file_names.pkl'.format(version_folder),'rb')))

        for i,file_name in enumerate(self.file_names):
            tables = []
            # if file is a .docx file and hasnt been seen yet
            if file_name not in finished_files:

                print('checking', file_name)
                file_str = '{}/{}'.format(folder,file_name)
                doc = docx.Document(file_str)

                # append the name to the first location if it exists
                if len(doc.tables) > 1: # are there at least 2 tables
                    # are there at least 2 rows and 3 columns
                    if len(doc.tables[1].rows) > 1 and len(doc.tables[1].columns) > 2:
                        beer_names.append(doc.tables[1].rows[1].cells[2].text)
                    else:
                        tables.append('NO NAME')
                        beer_names.append('NO NAME')
                else:
                    tables.append('NO NAME')
                    beer_names.append('NO NAME')

                # for each table in list of tables
                for idx, table in enumerate(doc.tables):
                    # if table is long enough
                    if len(table.rows)>2 and len(table.columns)>1:
                        # if it has a wanted comment field
                        if any([word in table.rows[2].cells[1].text.lower() for word in self.wanted_comments]):
                            # get the text from the table 2 after
                            table.rows[2].cells[1].text
                            df = process_table(doc.tables[idx+2])
                            tables.append(df)
            # check if it got read in correctly
            # name and three dfs and all tables same length
            if len(tables) == len(self.wanted_comments) and len(set([len(table) for table in tables])) == 1:
                beers.append(tables)
                ifiles.append(i)
                print('--> saved', file_name)

        categories = [f_name.split('/')[0] for f_name in self.file_names]
        self.category = np.array(categories)[ifiles]
        self.beer_tables = beers
        self.beer_names = np.array(beer_names)[ifiles]
        self.good_files = np.array(self.file_names)[ifiles]
        # self.ifiles = np.array(ifiles)
        #
        # # either update the files or save them fresh (cold start)
        # file_path_df = '{}/inter_files/docx_to_dfs.pkl'.format(version_folder)
        # self.beer_tables = self.Update_inter_file(file_path_df,beers)
        #
        # file_path_df = '{}/inter_files/file_names.pkl'.format(version_folder)
        # self.file_names = self.Update_inter_file(file_path_df,file_names)

        # return self.file_names, self.beer_tables

    def Get_text_df(self):
        '''
        turns the raw table data into a data frame of comments
        '''
        text_names = self.wanted_comments
        beer_names = []
        beer_text = []
        for beer in self.beer_tables:
            df_ = pd.DataFrame()
            for i, df in enumerate(beer):
                df_[text_names[i]] = df.iloc[:,-1]
            beer_text.append(df_)
        self.beer_text_dfs = beer_text
        return self.beer_text_dfs

    def Make_comment_dfs(self):
        '''
        makes a dataframe with 2 cols: beer_name and comment_text (all comments as a string)
        '''
        text = [' '.join([' '.join(col) for col in df.applymap(str).values.T]) for df in self.beer_text_dfs]

        # make beer_name col
        text_df = pd.DataFrame(list(zip(self.beer_names,text)))
        text_df.columns = ['beer_name','comment_text']
        # lower case the beer names
        text_df['beer_name'] = text_df['beer_name'].apply(lambda x: x.lower())
        # add underscores to beer names
        text_df['beer_name'] = text_df['beer_name'].apply(lambda x: re.sub(' ', '_', x))

        # make category col
        text_df['category'] = self.category
        # lower case the beer names
        text_df['category'] = text_df['category'].apply(lambda x: x.lower())
        # add underscores to beer names
        text_df['category'] = text_df['category'].apply(lambda x: re.sub(' ', '_', x))

        # combine duplicates
        text_df = text_df.groupby('beer_name').aggregate(lambda x: ' '.join(list(x)))

        text_df.reset_index(inplace = True)

        # get a list of off the category options
        text_df['category'] = text_df['category'].apply(lambda x: list(set(x.split())))

        # save / update the comment_text_df.pkl file
        text_df.to_pickle('{}/comment_text_df.pkl'.format(version_folder))
        self.text_df = text_df


    def Find_bad_format_files(self,folder):
        '''
        Returns the files that werent uploaded
        '''
        good_files = self.good_file
        all_files = os.listdir(folder)
        self.bad_files = [file_name for file_name in all_files if file_name not in good_files]
        return self.bad_files



if __name__ == '__main__':

    folder = 'files/doc_files' # .doc new beer files
    rnb = ReadNewBeers(cold_start = False)
    file_list = rnb.Get_files_from_folders('files/Brand Descriptions')
    rnb.Read_new_beer_files('files/Brand Descriptions',from_list = True)

    # files, beers = rnb.Read_new_beer_files('files/Brand Descriptions',from_list = True)
    text= rnb.Get_text_df()
    # rnb.Find_bad_format_files(folder)
    rnb.Make_comment_dfs()
