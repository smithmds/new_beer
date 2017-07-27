'''
Code written and developed by Jan Van Zeghbroeck
https://github.com/janvanzeghbroeck
'''


import docx
import os
import pandas as pd
import pickle
import re
import sys

# check the python version
if sys.version_info[0] == 3:
    version_folder = 'data3'
else:
    version_folder = 'data'# current_folder = os.getcwd()

def read_new_beer_files(folder,cold_start = False):

    '''
    reads all the tables from docx files and puts it into a list of lists of datafrems
    beers(file) --> beer --> tables --> table
    '''
    # file_name = 'Cherry Brown FTR Sep 16 15.docx'
    file_names = []
    beers = []
    if not cold_start:
        finished_files = list( pickle.load( open('../{}/inter_files/file_names.pkl'.format(version_folder),'rb')))
    else:
        finished_files = []

    for file_name in os.listdir(folder):
        if file_name in finished_files: #if we haven't seen this file yet
            pass
        else:
            print(file_name)
            file_str = '{}/{}'.format(folder,file_name)
            doc = docx.Document(file_str)

            tables = []
            # table_lst = doc.tables # all the tables as list
            for table in doc.tables: # doc.tables is type list
                n_rows = len(table.rows) #19
                # n_cols = len(table.columns)#6

                # creates a table
                table_ = []
                for i in range(n_rows):
                    row = table.rows[i]
                    row_ = []
                    for cell in row.cells:
                        row_.append(cell.text)
                    table_.append(row_)

                # creates data frame
                df = pd.DataFrame(table_)
                df.columns = df.iloc[0,:].values #rename columns
                df.drop(0,axis = 0,inplace = True) #delete column name row
                tables.append(df)

            if len(tables) == 22: #if not 22 tables than the file isn't formatted correctly
                beers.append(tables)
                file_names.append(file_name)

    if not cold_start:
        finished_beers_n_tables = list(pickle.load(open('../{}/inter_files/docx_to_dfs.pkl'.format(version_folder),'rb')))
        finished_file_names = list(pickle.load(open('../{}/inter_files/file_names.pkl'.format(version_folder),'rb')))
        beers = finished_beers_n_tables + beers
        file_names = finished_file_names + file_names

    pickle.dump( beers, open('../{}/inter_files/docx_to_dfs.pkl'.format(version_folder), "wb" ) )
    pickle.dump( file_names, open('../{}/inter_files/file_names.pkl'.format(version_folder), "wb" ) )

    # comments = df with all comments for FCAB
    comments = make_comment_dfs(beers)
    return comments, file_names, beers

'''
creates a dataframe for each beer with their comments in one string
to read: pd.read_pickle('data/inter_files/beer_batches.pkl')
'''


text_tables = [2,5,8,11] # indes of tables where we want the comments from
text_names = ['visual','aroma','flavor','body']

def get_text(beers):
    beers_text = []
    beer_names = []
    for beer in beers:
        df = pd.DataFrame()
        if len(beer) == 22: #28 dont have 22 tables
            for i in range(4):
                df_ = beer[text_tables[i]]
                text = df_.iloc[:,-1].values
                df[text_names[i]] = text
            beer_names.append(beer[1]['Name'].values[0])
            beers_text.append(df)
    return beers_text, beer_names


def make_comment_dfs(beers_n_tables = None):
    # beers(file) --> beer(176) --> tables(22 mostly) --> table
    # visual 2, aroma 5, flavor 8, mouthfeel 11
    if beers_n_tables is None:
        beers_n_tables = list(pickle.load(open('../{}/inter_files/docx_to_dfs.pkl'.format(version_folder),'rb')))

    beers, names = get_text(beers_n_tables)

    batch = pd.DataFrame()
    for col in text_names:
        new_col = []
        for beer in beers:
            new_col.append(' '.join(beer[col]))
        batch[col] = new_col
    batch['beer'] = names
    batch['beer'] = batch['beer'].apply(lambda x: x.lower())
    batch['beer'] = batch['beer'].apply(lambda x: re.sub(' ', '_', x))

    # remove duplicate
    batch = batch.groupby('beer').aggregate(lambda x: ' '.join(tuple(x)))

    batch.to_pickle('../{}/inter_files/beer_batches.pkl'.format(version_folder))
    return batch



def find_bad_format_files(folder):
    good_files = list( pickle.load( open('../{}/inter_files/file_names.pkl'.format(version_folder),'rb')))
    all_files = os.listdir(folder)
    bad_files = [file_name for file_name in all_files if file_name not in good_files]
    return bad_files




if __name__ == '__main__':
    folder = '../files/doc_files' # .doc new beer files
    comments, files, beers = read_new_beer_files(folder,cold_start = False)
    bad_files = find_bad_format_files(folder)
    # comments = make_comment_dfs()
